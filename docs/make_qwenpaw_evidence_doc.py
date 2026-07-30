# -*- coding: utf-8 -*-
"""Generate the preliminary-round QwenPaw development evidence document."""
from __future__ import annotations

import os

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from docstyle import (
    AZUL,
    CENTER,
    GREY,
    INK,
    TERRA,
    bullet,
    cjk,
    h1,
    h2,
    image,
    new_doc,
    page_break,
    para,
    runs,
    table,
)

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "assets")
QPAW = os.path.join(ASSETS, "qwenpaw")
OUT = os.path.join(HERE, "開發過程證明_QwenPaw_街知巷聞.docx")

doc = new_doc()

# Cover
p = doc.add_paragraph()
p.alignment = CENTER
p.paragraph_format.space_before = Pt(22)
r = p.add_run("QwenPaw 創新挑戰賽 · 初賽")
r.font.size = Pt(14)
r.font.color.rgb = AZUL
r.bold = True
cjk(r)

p = doc.add_paragraph()
p.alignment = CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run("開發過程證明")
r.font.size = Pt(28)
r.font.color.rgb = TERRA
r.bold = True
cjk(r)

para(doc, "街知巷聞 · EveryLane Macau", size=19, bold=True, align=CENTER, after=5)
para(
    doc,
    "QwenPaw Skill × MCP 工具鏈 × 澳門舊區智能導流",
    size=12,
    color=GREY,
    align=CENTER,
    after=18,
)

table(
    doc,
    [
        ["正式隊伍名稱", "愛拼才會贏"],
        ["參賽者 / 隊長", "施天益（SITINIEK）"],
        ["學號", "dc227126"],
        ["作品名稱", "街知巷聞 · EveryLane Macau"],
        ["QwenPaw 版本", "2.0.0（Windows 本地部署）"],
        ["證明日期", "2026 年 7 月 11 日"],
        ["公網實例（真實 Qwen）", "http://47.79.228.128/"],
        ["靜態備用演示", "https://mikey-si.github.io/everylane-macau/"],
    ],
    widths=[1.55, 4.8],
    header=False,
    firstcol=True,
)

image(
    doc,
    os.path.join(QPAW, "01_qwenpaw_console.png"),
    "QwenPaw 2.0.0 本地控制台已成功啟動（127.0.0.1:8088）",
    width=6.35,
)

# 1
h1(doc, "一、", "基礎部署：QwenPaw 已實際運行")
runs(
    doc,
    [
        ("本項目不是只在文字上「對標」QwenPaw，而是已在 Windows 本地完成 ", {}),
        ("QwenPaw 2.0.0", {"bold": True, "color": TERRA}),
        (" 部署，並建立獨立工作區、專用 Skill 與 MCP 客戶端。", {}),
    ],
)
table(
    doc,
    [
        ["節點", "實際結果", "可核對證據"],
        ["運行時", "QwenPaw FastAPI Console 正常返回 HTTP 200", "控制台截圖、啟動日誌"],
        ["智能體", "Default Agent；Plan Mode 已啟用", "agent.json / 控制台"],
        ["專用技能", "everylane-macau（local，Enabled）", "Skills 頁面截圖"],
        ["MCP 客戶端", "EveryLane Macau Tools（stdio，Enabled）", "MCP 頁面截圖"],
        ["模型接入", "支援 Token Plan OpenAI-compatible 協議", ".env.example / 連接診斷腳本"],
    ],
    widths=[1.15, 3.0, 2.2],
)

h2(doc, "安全處理")
bullet(doc, "API Key 只進入 QwenPaw 密鑰儲存或本機 .env；Git 忽略 .env。")
bullet(doc, "所有證明截圖避開 API Key 欄位；測試輸出只顯示「Key loaded: yes (masked)」。")
bullet(doc, "sk-sp Token Plan Key 與普通 DashScope Key 分流校驗，避免因 Base URL 不匹配出現 401。")

page_break(doc)
h1(doc, "二、", "真正接入：Skill + MCP 八項能力")
image(
    doc,
    os.path.join(QPAW, "02_everylane_skill_enabled.png"),
    "工作區專用 EveryLane Skill 已建立並啟用",
    width=6.35,
)
para(
    doc,
    "Skill 把「阿濠」人設、每日單一可步行片區、逐站開放核實、人流導流、"
    "低步行與預算控制寫成可重複執行的工作流程，而不是一次性 Prompt。",
)

image(
    doc,
    os.path.join(QPAW, "03_everylane_mcp_connected.png"),
    "EveryLane Macau Tools MCP 客戶端已連接（綠點 Enabled）",
    width=6.35,
)
table(
    doc,
    [
        ["MCP 工具", "任務", "可驗證輸出"],
        ["search_attractions", "按興趣 / 片區搜尋 70 POI", "真實 id、分類、標記、費用"],
        ["get_weather", "指定日期澳門天氣", "天氣、溫度、雨勢、建議"],
        ["check_opening", "逐站核實開放 / 休息日", "open、時段、失敗原因"],
        ["predict_crowd", "預測時段人流", "quiet / moderate / busy / packed"],
        ["find_local_gem", "熱點導流至老街 / 小店", "替代點、距離、步行分鐘"],
        ["compute_route", "計算步行順序", "每段距離、分鐘、總公里"],
        ["estimate_budget", "按人數估算費用", "逐項小計、總 MOP"],
        ["plan_macau_trip", "完整流程比較測試", "事件軌跡 + 結構化行程"],
    ],
    widths=[1.55, 2.25, 2.55],
)

page_break(doc)
h1(doc, "三、", "場景設計與智能體初步訓練 / 調優")
para(
    doc,
    "本項目的「初步訓練」不是擅自微調基礎模型權重，而是以場景集、Skill 指令、"
    "工具 Schema、確定性安全層與回歸測試反覆調優智能體行為。這種方式成本可控、"
    "可解釋，亦符合 QwenPaw 的技能與工具擴展機制。",
)
table(
    doc,
    [
        ["訓練 / 調優場景", "首輪問題", "調優後行為"],
        ["鄭家大屋・星期三", "模型可能照排休息景點", "check_opening → 失敗恢復 → 同區替代"],
        ["大三巴・週末正午", "只描述「很擠」", "predict_crowd → 錯峰 → find_local_gem 導流"],
        ["爸媽 / 長者", "路線站點多、步行遠", "少行路上限 3.6 km，自動刪除最遠站"],
        ["三日兩夜", "同一天跨島、跨日重複", "半島 / 氹仔 / 路環分日，used_ids 去重"],
        ["低預算", "推薦收費點後只提示超支", "移除昂貴非必要點，再次 estimate_budget"],
        ["模型提前交卷", "未用齊工具也可 submit", "缺工具 / 未查開放即拒絕提交並繼續 ReAct"],
        ["五語言", "葡語 / 日語動態結果回退英文", "固定 UI、工具軌跡、行程、核對、貼士全本地化"],
    ],
    widths=[1.4, 2.25, 2.7],
)

h2(doc, "代表性 ReAct 軌跡")
table(
    doc,
    [
        ["步驟", "QwenPaw / MCP 行動", "觀察與決策"],
        ["1", "get_weather(星期三)", "按雨勢調整室內外比例"],
        ["2", "search_attractions(history, prefer_local)", "取得同區候選與本地小店"],
        ["3", "check_opening(mandarin_house)", "逢週三休息 → 觸發恢復"],
        ["4", "check_opening(替代點)", "確認有開才納入"],
        ["5", "predict_crowd(熱門地標)", "busy / packed → 錯峰"],
        ["6", "find_local_gem(熱門地標)", "導流至附近舊街"],
        ["7", "compute_route + estimate_budget", "核對距離與預算後提交"],
    ],
    widths=[0.6, 3.0, 2.75],
)

page_break(doc)
h1(doc, "四、", "五輪測試生命周期與量化結果")
table(
    doc,
    [
        ["輪次", "測試 → 發現 → 修復", "最終結果"],
        ["1 後端正確性", "解析、開放、路線、多日、圖片、預算、人流", "531 PASS / 0 FAIL / 0 WARN"],
        ["2 前端體驗", "五語言、自由輸入、圖片、手機、列印、SSE", "99 PASS / 0 FAIL"],
        ["3 健壯安全", "注入、超長、穿越、並發、安全頭、Qwen 防線", "64 PASS / 0 FAIL"],
        ["4 QwenPaw", "版本、Skill、MCP 協議、8 工具、失敗恢復", "MCP E2E PASS"],
        ["5 最終回歸", "後端 + 前端 + API + MCP + 文檔一致性", "見最終 QA 報告"],
    ],
    widths=[1.25, 3.5, 1.6],
)
bullet(doc, "70 / 70 個 POI 有真實圖片；25 個舊區點、20 間本地小店。")
bullet(doc, "12 個核心需求場景 + 9 種異常輸入 + 12 條並發流全部完成。")
bullet(doc, "桌面 1440px、手機 375px、GitHub Pages、FastAPI SSE 四種模式均測試。")
bullet(doc, "QwenPaw MCP 以真實 ClientSession 完成 list_tools 與 call_tool，而非只測 Python 函數。")

h2(doc, "核心測試檔")
table(
    doc,
    [
        ["檔案", "用途"],
        ["qa/test_backend.py", "知識庫、解析、12 場景、時間 / 路線 / 預算不變量"],
        ["qa/test_frontend.py", "Playwright 五語言、桌面 / 手機、靜態 / SSE、列印"],
        ["qa/test_api.py", "安全、邊界輸入、12 並發、目錄穿越、Qwen 安全層"],
        ["qa/test_qwenpaw_mcp.py", "MCP 協議、8 工具、休息日恢復與完整規劃"],
    ],
    widths=[2.1, 4.25],
)

page_break(doc)
h1(doc, "五、", "架構與可重現性")
table(
    doc,
    [
        ["層級", "實作"],
        ["QwenPaw 運行時", "QwenPaw 2.0.0 Console / Default Agent / Plan Mode"],
        ["模型", "Token Plan：qwen3.7-plus（配置就緒；Key 由隊長在密鑰頁安全貼入）"],
        ["行為層", "EveryLane Skill：阿濠人設、流程、約束、輸出規格"],
        ["工具層", "stdio MCP：7 個細粒度工具 + 1 個整體比較工具"],
        ["領域層", "70 POI 澳門知識庫、開放、人流、地理、預算"],
        ["產品層", "FastAPI SSE + 五語言 Web + Leaflet 地圖 + PDF"],
        ["驗證層", "自动化 QA + 五輪生命周期报告"],
    ],
    widths=[1.45, 4.9],
    header=False,
    firstcol=True,
)
para(
    doc,
    "可重現步驟：安裝 QwenPaw → 在模型頁加入主辦方 Token Plan Key（截圖須打碼）"
    "→ 啟用 EveryLane Skill 與 MCP → 輸入指定情境 → 查看逐步工具調用與最終行程。"
)
para(
    doc,
    "詳細配置、MCP 路徑、驗收提示詞與截圖清單見 qwenpaw/README.md。",
    color=GREY,
)

h1(doc, "六、", "結論")
runs(
    doc,
    [
        ("目前已具備初賽要求的 ", {}),
        ("QwenPaw 基礎部署、場景設計、智能體調優與開發過程證明", {"bold": True, "color": TERRA}),
        ("。核心競爭力不是一般旅遊問答，而是可驗證的多工具 ReAct 流程，以及把熱門客流導向舊區小店的可量化社會與商業價值。", {}),
    ],
)
para(doc, "— 完 —", align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, before=12)

doc.save(OUT)
print("Saved:", OUT)
