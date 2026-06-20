# -*- coding: utf-8 -*-
"""Generate the competition Concept Design Document (概念計劃書) as a .docx,
following the official sample structure but elevated for 街知巷聞 / EveryLane Macau."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "概念計劃書_街知巷聞_EveryLaneMacau.docx")
ASSET_DIR = os.path.join(HERE, "assets")
PAGES_URL = "https://mikey-si.github.io/everylane-macau/"

TERRA = RGBColor(0xBE, 0x4A, 0x3A)
AZUL = RGBColor(0x2C, 0x5E, 0x86)
INK = RGBColor(0x2B, 0x21, 0x18)
GOLD = RGBColor(0x9C, 0x6E, 0x1E)
GREY = RGBColor(0x6b, 0x60, 0x55)

BODY_FONT = "Microsoft JhengHei"
SERIF_FONT = "Microsoft JhengHei"

doc = Document()
for sec in doc.sections:
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Inches(0.62)
    sec.bottom_margin = Inches(0.62)
    sec.left_margin = Inches(0.68)
    sec.right_margin = Inches(0.68)
# base style
st = doc.styles["Normal"]
st.font.name = BODY_FONT
st.font.size = Pt(10.5)
st.font.color.rgb = INK
st._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
st.paragraph_format.space_after = Pt(5)
st.paragraph_format.line_spacing = 1.28


def _cjk(run, font=BODY_FONT):
    run.font.name = font
    r = run._element
    rpr = r.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:eastAsia"), font)
    rf.set(qn("w:ascii"), font)
    rf.set(qn("w:hAnsi"), font)


def para(text="", size=10.5, color=INK, bold=False, align=None, before=0, after=5, italic=False, font=BODY_FONT):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = color
        _cjk(r, font)
    return p


def runs(p, parts):
    """parts: list of (text, {opts})"""
    for text, o in parts:
        r = p.add_run(text)
        r.bold = o.get("bold", False)
        r.italic = o.get("italic", False)
        r.font.size = Pt(o.get("size", 10.5))
        r.font.color.rgb = o.get("color", INK)
        _cjk(r, o.get("font", BODY_FONT))
    return p


def h1(no, text):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{no}　{text}")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = TERRA
    _cjk(r, SERIF_FONT)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "E0CBA8")
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("▎" + text)
    r.bold = True
    r.font.size = Pt(11.8)
    r.font.color.rgb = AZUL
    _cjk(r, SERIF_FONT)
    return p


def bullet(text, bold_head=None):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.26)
    p.paragraph_format.space_after = Pt(3)
    dot = p.add_run("· ")
    dot.font.color.rgb = GOLD; dot.bold = True; _cjk(dot)
    if bold_head:
        r = p.add_run(bold_head)
        r.bold = True; r.font.color.rgb = INK; r.font.size = Pt(10.5); _cjk(r)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5); r2.font.color.rgb = INK; _cjk(r2)
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def cell_margin(cell, top=70, start=90, bottom=70, end=90):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def no_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))


def set_cell(cell, text, bold=False, color=INK, size=10, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if white else color
    _cjk(r)


def table(rows, widths=None, header=False, head_fill="2C5E86"):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = False
    for ri, row in enumerate(rows):
        no_row_split(t.rows[ri])
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell_margin(cell)
            if header and ri == 0:
                shade(cell, head_fill); set_cell(cell, val, bold=True, white=True, size=10)
            else:
                if ci == 0 and not header:
                    shade(cell, "F4EAD8"); set_cell(cell, val, bold=True, size=10)
                else:
                    set_cell(cell, val, size=10)
    if widths:
        for ci, w in enumerate(widths):
            for r in t.rows:
                r.cells[ci].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_image(name, caption, width=6.35, before=4, after=8):
    path = os.path.join(ASSET_DIR, name)
    if not os.path.exists(path):
        return None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(path, width=Inches(width))
    cap = para("圖：" + caption, size=8.8, color=GREY, italic=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, after=after)
    cap.paragraph_format.keep_together = True
    return p


# ====================== TITLE ======================
para("「千模百煉」AI 開發者系列之學生競賽", 11, GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("參賽作品概念設計書", 12, AZUL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("街知巷聞"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = TERRA; _cjk(r, SERIF_FONT)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run("EveryLane Macau"); r.font.size = Pt(15); r.font.color.rgb = GREY
p2.paragraph_format.space_after = Pt(2)
para("澳門深度遊 AI 智能體 —— 唔止大三巴，帶你行勻澳門每一條老街", 12, INK, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=14)

table([
    ["參賽方向", "澳門文旅專題（兼及舊區活化）"],
    ["作品形式", "Web 應用（可運行網站）＋ 基於 Qwen / QwenPaw 的 ReAct AI 智能體"],
    ["線上展示", PAGES_URL],
    ["團隊名稱", "街知巷聞工作室"],
    ["參賽者 / 隊長", "SITINIEK（學號 dc227126）"],
    ["指導教師", "（待定）"],
    ["日期", "2026 年 6 月"],
], widths=[1.6, 4.7])

add_image("product_hero.png", "網站首頁視覺：以澳門舊城色系呈現「文旅 × 舊區活化」定位", width=6.35, after=10)

# ====================== 一 ======================
h1("一、", "我們想做什麼")
para("我們想用 Qwen 千問大模型與 QwenPaw 智能體框架，打造一個真正「會做任務」的澳門深度遊 AI 智能體"
     "——「街知巷聞」（人設：本地老街坊「阿濠」）。")
runs(doc.add_paragraph(), [
    ("用戶只要用一句自然語言講低需求（例如：「我下星期三想帶爸媽嚟玩一日」或「幫我安排澳門三日兩夜」），"
     "阿濠就會自主完成以下任務，並輸出一份", {}),
    ("可驗證", {"bold": True, "color": TERRA}),
    ("的行程：", {}),
])
bullet("查當日天氣、搜尋合適景點；")
bullet("逐一核實每個景點當日是否開放（避開休息日）；")
bullet("預測各景點人流，並把遊客由逼爆的熱點，智能導流到舊區老街與本地老字號；")
bullet("計算順路的步行路線、估算人均預算，遇到限制衝突會自動改線。")
para("一句話：阿濠唔係一個只會聊天的問答機械人，而係一個識得規劃、調用工具、多步執行、"
     "甚至失敗自動恢復的智能體 —— 把「揾景點、查資料、砌行程」呢啲繁瑣工作一手包辦。", color=INK, bold=False)

# ====================== 二 ======================
h1("二、", "為什麼做這個（痛點與價值）")
para("澳門旅遊業高度集中：遊客逼爆大三巴、議事亭、威尼斯人，而福隆新街、十月初五街、下環等"
     "舊區老街與街坊老字號卻日漸凋零。呢個現象同時帶嚟三個問題：")
bullet("遊客體驗下降：熱點人多、排隊、千篇一律，錯過真正有文化底蘊的澳門；", "對遊客　")
bullet("舊區小店客流不足、文化記憶流失，活化困難；", "對社區　")
bullet("旅遊承載失衡、過度集中，與「世界旅遊休閒中心」的高質定位有落差。", "對城市　")
para("我們相信 AI 智能體可以做一件有商業價值又有社會意義的事：把流量「重新分配」"
     "—— 用智能導流，將遊客有策略地引入舊區老街與本地小店，做到遊客、小店、城市三方共贏。")

# ====================== 三 ======================
h1("三、", "怎麼做（技術方案）")
para("作品為一個可直接運行的網站，後端以 FastAPI 承載一個 ReAct 智能體，"
     "對標 QwenPaw「FastAPI 運行時 + ReAct 核心 + 工具 / Toolkit」的五層架構。", )
h2("步驟 1：建立澳門景點知識庫（真實、可驗證）")
bullet("以程式抓取 Wikipedia / Wikimedia Commons 的真實景點坐標與相片（公開授權）；")
bullet("人手整理開放時間、休息日、費用、人流特徵、舊區 / 本地小店標記，共 70 個景點。")
h2("步驟 2：以 Qwen + QwenPaw 構建 ReAct 智能體")
bullet("透過百煉（DashScope）OpenAI 相容接口接入 Qwen 千問模型，驅動 ReAct「思考-行動-觀察」迴圈；")
bullet("為智能體配置 7 種工具，統一以 function-calling schema 註冊（對標 QwenPaw Toolkit）：")
para("　　搜尋景點 · 查天氣 · 核實開放時間 · 預測人流 · 舊區導流 · 規劃步行路線 · 估算預算", size=10, color=AZUL, bold=True)
bullet("設定阿濠人設 Prompt（粵語、親切、本地街坊味），並要求所有結論基於工具返回的真實數據。")
h2("步驟 3：實時可視化 + 測試優化")
bullet("前端以 Server-Sent Events 實時展示智能體的規劃、工具調用、改線過程，並以地圖 + 時間軸呈現行程；")
bullet("反覆測試多種情境（不同興趣 / 人數 / 預算 / 日期），校正路線合理性與失敗恢復邏輯。")
para("所需資源：一台個人電腦 + 大會提供的百煉代金券。系統亦內建「離線示範引擎」，"
     "未有金鑰時仍能調用同一套真實工具完整運行，方便開發與路演。", color=GREY)

# ====================== 四 ======================
doc.add_page_break()
h1("四、", "智能體工作示例（真實運行軌跡）")
para("以下為一次真實運行的智能體軌跡（用戶：「我想去鄭家大屋同附近嘅歷史老街，星期三去」）：")
table([
    ["步驟", "智能體行動", "結果 / 決策"],
    ["規劃", "解析需求：日期=星期三、興趣=歷史/老街、片區=澳門半島", "鎖定可步行的歷史城區，避免跨島"],
    ["工具", "get_weather(週三)", "短暫陣雨 → 多安排室內景點"],
    ["工具", "search_attractions(歷史/老街, prefer_local)", "取得草堆街、福隆新街、鄭家大屋等候選"],
    ["工具", "check_opening(鄭家大屋, 星期三)", "✗ 逢週三休息"],
    ["失敗恢復", "於同片區改揀有開、同類的替代點", "自動改為「戀愛巷」，行程不受影響"],
    ["工具", "predict_crowd(大三巴, 13:00)", "極擁擠 → 觸發導流"],
    ["導流", "find_local_gem(大三巴)", "加插「草堆街與爛鬼樓」（步行 6 分）分流"],
    ["工具", "compute_route + estimate_budget", "全程步行 1.2 km、人均 MOP 0，輸出行程"],
], widths=[0.9, 3.2, 2.4], header=True)
para("最終輸出包含：地圖路線、逐站時間軸（到/離時間、人流、費用、步行距離）、"
     "以及「任務完成核對」面板（預算、步行、開放、避開人潮、帶旺舊區逐項驗證）。", color=GREY)
add_image("route_map.png", "行程路線地圖：真實坐標、步行順序與熱點導流結果可視化", width=6.35, after=8)

# ====================== 五 ======================
h1("五、", "創意亮點")
bullet("把 AI 由「問答」升級為「做任務」：規劃 + 工具調用 + 多步執行 + 失敗自動恢復，貼合賽事對智能體能力的要求。", "真·智能體　")
bullet("獨家定位「導流引擎」：主動將遊客由人潮熱點分流到舊區老街與本地老字號，文旅 × 舊區活化雙賽道共振。", "舊區導流　")
bullet("每一項結論（開放時間、人流、步行距離、預算）都列明、可核對，唔係空泛建議。", "結果可驗證　")
bullet("實時流式展示智能體「諗嘢 / 用工具 / 改線」全過程，評審一眼睇懂其能力。", "過程可視化　")
bullet("阿濠用澳門粵語口吻講古、講路線，仲支援普通話 / 英文 / 葡文 / 日文多語言。", "在地 + 多語　")

# ====================== 六 ======================
doc.add_page_break()
h1("六、", "計劃時間表")
table([
    ["時間", "任務"],
    ["7 月上旬", "完成 70 個景點知識庫，補全開放時間 / 人流 / 相片；接入百煉 Qwen"],
    ["7 月中旬", "完善 ReAct 工具鏈、多日分區行程與失敗恢復策略；多情境測試與路線校正"],
    ["7 月下旬", "前端體驗打磨（多日總覽、地圖、時間軸、多語言）；加入收藏 / 匯出 PDF"],
    ["8 月上旬", "邀請 20+ 位同學與遊客試用，收集反饋；接入真實天氣 API"],
    ["8 月中旬", "優化迭代：人流模型校正、商戶端原型（精選商戶 / 引流統計）"],
    ["8 月下旬", "撰寫實踐文章、製作路演 Demo 與展示材料"],
], widths=[1.4, 4.9], header=True, head_fill="BE4A3A")

# ====================== 七 ======================
h1("七、", "預期效果與價值")
bullet("覆蓋 70 個澳門景點，當中 25 個為舊區老街 / 活化片區、20 間為本地小店；")
bullet("智能體能就任意合理需求，輸出一日或 2–5 日多日深度遊，並確保每日鎖定可步行片區；")
bullet("每份行程平均納入 ≥ 3 個舊區老街 / 本地小店，量化「導流」成效；")
bullet("試用者中 80% 認同「比自己上網查更慳時間、更有在地味」。")
para("社會與商業價值：本作品不只是 AI 應用，更是一個可持續的「舊區導流引擎」——"
     "為小店帶客流、為遊客帶體驗、為城市帶平衡，具清晰的 B 端變現路徑"
     "（商戶精選訂閱 / 引流分成、酒店與旅行社 API、文旅數據儀表板）。", bold=False)

# ====================== 八 ======================
h1("八、", "AI 倫理考量")
bullet("行程由 AI 生成，介面明確標示；人流與天氣為估算值，提醒用戶以現場為準。", "透明　")
bullet("景點開放時間、坐標等關鍵事實一律基於知識庫，模型不得杜撰；超出知識庫範圍不亂作答。", "準確　")
bullet("刻意把流量導向資源較弱的舊區小店，促進社區可持續，而非加劇過度集中。", "公平與在地關懷　")
bullet("不收集個人身分資料、無需登入即可使用；景點相片均採用公開授權（Wikimedia Commons）。", "私隱與版權　")

# ====================== 九 ======================
h1("九、", "分工與技術說明")
para("本作品由參賽者 SITINIEK（學號 dc227126）獨立完成，涵蓋以下範疇：")
table([
    ["範疇", "內容"],
    ["產品 / 創意", "舊區導流定位、情境設計、商業模式"],
    ["資料工程", "景點知識庫建構、Wikipedia/Commons 資料抓取與校正"],
    ["智能體 / 後端", "ReAct 迴圈、7 種工具、Qwen 接入、FastAPI 與 SSE"],
    ["前端 / 設計", "互動網站、地圖與時間軸、智能體過程可視化"],
], widths=[1.8, 4.5], header=True)
para("我們相信，Qwen 與 QwenPaw 讓「諗到」就能「做到」——"
     "把對澳門舊區的關懷，變成一個真正幫到遊客同街坊的智能體。", bold=True, color=TERRA, before=6)
para("— 完 —", align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, before=10)

doc.save(OUT)
print("Saved:", OUT)
