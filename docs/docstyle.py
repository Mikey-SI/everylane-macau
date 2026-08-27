# -*- coding: utf-8 -*-
"""Shared docx styling helpers for the competition documents."""
from docx import Document
import os

from docx.shared import Pt, RGBColor, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TERRA = RGBColor(0xBE, 0x4A, 0x3A)
AZUL = RGBColor(0x2C, 0x5E, 0x86)
INK = RGBColor(0x2B, 0x21, 0x18)
GOLD = RGBColor(0x9C, 0x6E, 0x1E)
GREY = RGBColor(0x6B, 0x60, 0x55)
CENTER = WD_ALIGN_PARAGRAPH.CENTER

BODY_FONT = "Microsoft JhengHei"
SERIF_FONT = "Microsoft JhengHei"
MONO = "Consolas"


def new_doc():
    doc = Document()
    # A4 with slightly wider printable area. This keeps tables from wrapping
    # awkwardly and gives screenshots enough width to look crisp.
    for sec in doc.sections:
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.top_margin = Inches(0.62)
        sec.bottom_margin = Inches(0.62)
        sec.left_margin = Inches(0.68)
        sec.right_margin = Inches(0.68)
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = Pt(10.5)
    st.font.color.rgb = INK
    st._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    st.paragraph_format.space_after = Pt(5)
    st.paragraph_format.line_spacing = 1.3
    return doc


def cjk(run, font=BODY_FONT):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:eastAsia"), font); rf.set(qn("w:ascii"), font); rf.set(qn("w:hAnsi"), font)


def para(doc, text="", size=10.5, color=INK, bold=False, align=None, before=0, after=5, italic=False, font=BODY_FONT):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic; r.font.size = Pt(size); r.font.color.rgb = color
        cjk(r, font)
    return p


def runs(doc, parts, after=5, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    for text, o in parts:
        r = p.add_run(text)
        r.bold = o.get("bold", False); r.italic = o.get("italic", False)
        r.font.size = Pt(o.get("size", 10.5)); r.font.color.rgb = o.get("color", INK)
        cjk(r, o.get("font", BODY_FONT))
    return p


def h1(doc, no, text):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{no}　{text}" if no else text)
    r.bold = True; r.font.size = Pt(15); r.font.color.rgb = TERRA; cjk(r, SERIF_FONT)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "E0CBA8")
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run("▎" + text)
    r.bold = True; r.font.size = Pt(11.8); r.font.color.rgb = AZUL; cjk(r, SERIF_FONT)
    return p


def bullet(doc, text, head=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.26); p.paragraph_format.space_after = Pt(3)
    dot = p.add_run("· "); dot.font.color.rgb = GOLD; dot.bold = True; cjk(dot)
    if head:
        r = p.add_run(head); r.bold = True; r.font.size = Pt(10.5); cjk(r)
    r2 = p.add_run(text); r2.font.size = Pt(10.5); r2.font.color.rgb = INK; cjk(r2)
    return p


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _cell_margin(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def _no_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))


def _set(cell, text, bold=False, color=INK, size=10, white=False, mono=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.keep_together = True
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if white else color
    cjk(r, MONO if mono else BODY_FONT)


def table(doc, rows, widths=None, header=True, head_fill="2C5E86", firstcol=False):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = False
    for ri, row in enumerate(rows):
        _no_row_split(t.rows[ri])
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            _cell_margin(cell)
            if header and ri == 0:
                _shade(cell, head_fill); _set(cell, val, bold=True, white=True)
            elif (firstcol or not header) and ci == 0:
                _shade(cell, "F4EAD8"); _set(cell, val, bold=True)
            else:
                _set(cell, val)
    if widths:
        for ci, w in enumerate(widths):
            for r in t.rows:
                r.cells[ci].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def page_break(doc):
    doc.add_page_break()


def _set_run_border(p, color="BE4A3A", sz="18"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _add_field(paragraph, instr):
    run = paragraph.add_run()
    r = run._r
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText"); text.set(qn("xml:space"), "preserve"); text.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    r.append(begin); r.append(text); r.append(sep); r.append(end)


def header_footer(doc, header_left="街知巷聞 · EveryLane Macau", header_right="複賽說明文檔",
                  footer_left="http://47.79.228.128/"):
    """Cover has no running header; subsequent pages get a thin header + page numbers."""
    for sec in doc.sections:
        sec.different_first_page_header_footer = True
        sec.header_distance = Mm(8)
        sec.footer_distance = Mm(8)

        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(2)
        r = hp.add_run(header_left)
        r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = TERRA; cjk(r)
        r2 = hp.add_run("　　" + header_right)
        r2.font.size = Pt(8.5); r2.font.color.rgb = GREY; cjk(r2)
        _set_run_border(hp, "E0CBA8", "8")

        fp = sec.footer.paragraphs[0]
        fp.alignment = CENTER
        fp.paragraph_format.space_before = Pt(2)
        r = fp.add_run(footer_left + "  ·  ")
        r.font.size = Pt(8); r.font.color.rgb = GREY; cjk(r)
        r = fp.add_run("第 ")
        r.font.size = Pt(8); r.font.color.rgb = GREY; cjk(r)
        _add_field(fp, " PAGE ")
        r = fp.add_run(" / ")
        r.font.size = Pt(8); r.font.color.rgb = GREY
        _add_field(fp, " NUMPAGES ")
        r = fp.add_run(" 頁")
        r.font.size = Pt(8); r.font.color.rgb = GREY; cjk(r)


def callout(doc, title, body):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    _shade(cell, "F4EAD8")
    _cell_margin(cell, top=90, start=140, bottom=90, end=140)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = TERRA; cjk(r)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10); r2.font.color.rgb = INK; cjk(r2)
    for border in ("top", "left", "bottom", "right"):
        tc = cell._tc.get_or_add_tcPr()
        # keep a gold-tinted grid via table style; extra shade is enough
        _ = border
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def _clear_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def image(doc, filename, caption, width=6.35, before=4, after=8):
    """Add a centred image with a compact caption if the file exists."""
    if not os.path.exists(filename):
        return None
    p = doc.add_paragraph()
    p.alignment = CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    r = p.add_run()
    r.add_picture(filename, width=Inches(width))
    cap = para(doc, "圖：" + caption, size=8.6, color=GREY, italic=True, align=CENTER, after=after)
    cap.paragraph_format.keep_together = True
    return p


def two_images(doc, left, left_cap, right, right_cap, width=3.12):
    """Side-by-side screenshots with captions — skips missing files."""
    files = [(left, left_cap), (right, right_cap)]
    files = [(f, c) for f, c in files if os.path.exists(f)]
    if not files:
        return None
    if len(files) == 1:
        return image(doc, files[0][0], files[0][1], width=6.3)
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for ci, (fn, cap) in enumerate(files):
        img_cell = t.cell(0, ci)
        cap_cell = t.cell(1, ci)
        img_cell.width = Inches(width + 0.12)
        cap_cell.width = Inches(width + 0.12)
        _clear_cell_borders(img_cell)
        _clear_cell_borders(cap_cell)
        _cell_margin(img_cell, top=40, start=40, bottom=20, end=40)
        _cell_margin(cap_cell, top=0, start=40, bottom=60, end=40)
        img_cell.text = ""
        p = img_cell.paragraphs[0]
        p.alignment = CENTER
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(fn, width=Inches(width))
        cap_cell.text = ""
        p2 = cap_cell.paragraphs[0]
        p2.alignment = CENTER
        r = p2.add_run("圖：" + cap)
        r.italic = True; r.font.size = Pt(8.2); r.font.color.rgb = GREY; cjk(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def qr_with_caption(doc, filename, caption, width=1.28):
    if not os.path.exists(filename):
        return None
    p = doc.add_paragraph()
    p.alignment = CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(filename, width=Inches(width))
    return para(doc, caption, size=8.4, color=GREY, italic=True, align=CENTER, after=8)
